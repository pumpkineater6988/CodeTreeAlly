from flask import Flask, render_template, Blueprint, request, redirect, url_for, session, send_from_directory
from flask_sqlalchemy import SQLAlchemy
import os
from PIL import Image as PILImage
import uuid
import fitz  # PyMuPDF
import docx
import pandas as pd
from io import BytesIO
from docx import Document
import openai
from dotenv import load_dotenv
import requests
import numpy as np
from markupsafe import Markup

load_dotenv()
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')
GEMINI_API_KEY_CHAT = os.environ.get('GEMINI_API_KEY_CHAT')
GEMINI_API_KEY_IMAGE = os.environ.get('GEMINI_API_KEY_IMAGE')
GEMINI_API_URL = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro-latest:generateContent'

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///genai_app.db'
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['SESSION_TYPE'] = 'filesystem'
db = SQLAlchemy(app)

# Ensure folders exist
os.makedirs('static', exist_ok=True)
os.makedirs('templates', exist_ok=True)

# Define blueprints (text_to_image bp REMOVED)
document_assistant_bp = Blueprint('document_assistant', __name__, url_prefix='/document-assistant')
developer_copilot_bp = Blueprint('developer_copilot', __name__, url_prefix='/developer-copilot')
internal_chatbot_bp = Blueprint('internal_chatbot', __name__, url_prefix='/internal-chatbot')
chat_history_bp = Blueprint('chat_history', __name__, url_prefix='/chat-history')
prompt_optimizer_bp = Blueprint('prompt_optimizer', __name__, url_prefix='/prompt-optimizer')

@document_assistant_bp.route('/', methods=['GET', 'POST'])
def document_assistant_home():
    result = None
    error = None
    file_text = None
    if request.method == 'POST':
        if 'document' not in request.files or not request.files['document'].filename:
            error = 'Document file is required.'
        else:
            doc_file = request.files['document']
            filename = doc_file.filename.lower() if doc_file.filename else ''
            query_type = request.form.get('query_type')
            query = request.form.get('query')
            try:
                text = ''
                if filename.endswith('.pdf'):
                    try:
                        import fitz
                        with fitz.open(stream=doc_file.read(), filetype='pdf') as pdf:
                            text = '\n'.join([page.get_text() for page in pdf])
                    except ImportError:
                        text = "PyMuPDF (fitz) not installed. Please install: pip install PyMuPDF"
                    except Exception as e:
                        text = f"Error reading PDF: {str(e)}"
                elif filename.endswith('.docx'):
                    try:
                        docx_bytes = doc_file.read()
                        from docx import Document
                        doc = Document(BytesIO(docx_bytes))
                        text = '\n'.join([p.text for p in doc.paragraphs])
                    except ImportError:
                        text = "python-docx not installed. Please install: pip install python-docx"
                    except Exception as e:
                        text = f"Error reading DOCX: {str(e)}"
                elif filename.endswith('.txt'):
                    text = doc_file.read().decode('utf-8', errors='ignore')
                elif filename.endswith('.csv'):
                    try:
                        csv_bytes = doc_file.read()
                        import pandas as pd
                        df = pd.read_csv(BytesIO(csv_bytes))
                        text = df.to_string()
                    except ImportError:
                        text = "pandas not installed. Please install: pip install pandas"
                    except Exception as e:
                        text = f"Error reading CSV: {str(e)}"
                elif filename.endswith('.xlsx'):
                    try:
                        xlsx_bytes = doc_file.read()
                        import pandas as pd
                        df = pd.read_excel(BytesIO(xlsx_bytes))
                        text = df.to_string()
                    except ImportError:
                        text = "pandas not installed. Please install: pip install pandas openpyxl"
                    except Exception as e:
                        text = f"Error reading XLSX: {str(e)}"
                elif filename.endswith(('.png', '.jpg', '.jpeg')):
                    # OCR for images containing text/tables
                    try:
                        import cv2
                        import pytesseract
                        from PIL import Image
                        import numpy as np
                        
                        # Read image bytes and convert to PIL Image
                        image_bytes = doc_file.read()
                        pil_image = Image.open(BytesIO(image_bytes))
                        
                        # Convert to OpenCV format
                        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
                        
                        # Preprocess image for better OCR
                        gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
                        # Apply threshold to get black text on white background
                        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        
                        # Extract text using Tesseract OCR
                        text = pytesseract.image_to_string(thresh, config='--psm 6')
                        
                        if not text.strip():
                            # Try with different PSM mode for tables
                            text = pytesseract.image_to_string(thresh, config='--psm 3')
                        
                        if not text.strip():
                            text = "No text could be extracted from this image."
                    except ImportError:
                        text = "OCR libraries not installed. Please install opencv-python, pytesseract, and pillow."
                    except Exception as e:
                        text = f"OCR error: {str(e)}"
                else:
                    error = 'Unsupported file type. Supported: PDF, DOCX, TXT, CSV, XLSX, PNG, JPG, JPEG'
                file_text = text
                if not error:
                    # Try Gemini API first if available
                    if GEMINI_API_KEY_CHAT and query_type in ['summary', 'extract', 'context']:
                        try:
                            prompt = f"Document content:\n{file_text}\n\nUser query: {query or query_type}"
                            headers = {'Content-Type': 'application/json'}
                            params = {'key': GEMINI_API_KEY_CHAT}
                            data = {"contents": [{"parts": [{"text": prompt}]}]}
                            response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data)
                            if response.status_code == 200:
                                gemini_data = response.json()
                                result = gemini_data['candidates'][0]['content']['parts'][0]['text']
                            else:
                                # Fall back to basic processing if Gemini fails
                                result = process_document_basic(file_text, query_type, query)
                        except Exception as e:
                            # Fall back to basic processing if Gemini fails
                            result = process_document_basic(file_text, query_type, query)
                    else:
                        # Basic document processing without AI
                        result = process_document_basic(file_text, query_type, query)
            except Exception as e:
                error = f'Error processing document: {str(e)}'
    return render_template('document_assistant.html', result=result, error=error)

@developer_copilot_bp.route('/', methods=['GET', 'POST'])
def developer_copilot_home():
    suggestion = None
    error = None
    if request.method == 'POST':
        code = request.form.get('code')
        if not code:
            error = 'Please paste your code or error message.'
        else:
            if GEMINI_API_KEY_CHAT:
                prompt = (
                    "Explain the following error or code in simple terms, why it happens, and provide a fixed code snippet if possible. "
                    "Format the fix as a code block (markdown triple backticks) so it is copyable.\n\n" + code
                )
                headers = {'Content-Type': 'application/json'}
                params = {'key': GEMINI_API_KEY_CHAT}
                data = {"contents": [{"parts": [{"text": prompt}]}]}
                try:
                    response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data)
                    if response.status_code == 200:
                        gemini_data = response.json()
                        suggestion = gemini_data['candidates'][0]['content']['parts'][0]['text']
                    else:
                        suggestion = f"[Gemini API error: {response.text}]"
                except Exception as e:
                    suggestion = f"[Gemini error: {e}]"
            else:
                suggestion = 'Check your syntax and imports. If the error persists, review the stack trace and documentation.'
    return render_template('developer_copilot.html', suggestion=suggestion, error=error)

OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')
if OPENAI_API_KEY:
    openai.api_key = OPENAI_API_KEY

@internal_chatbot_bp.route('/', methods=['GET', 'POST'])
def internal_chatbot_home():
    if 'chat_history' not in session:
        session['chat_history'] = []
    chat_history = session['chat_history']
    error = None
    if request.method == 'POST':
        user_text = request.form.get('question')
        file_name = None
        file_text = None
        # Handle file upload and extract text if possible
        if 'file' in request.files and request.files['file'].filename:
            file = request.files['file']
            file_name = file.filename
            save_path = os.path.join('static', 'uploads')
            os.makedirs(save_path, exist_ok=True)
            file_path = os.path.join(save_path, file_name)
            file.save(file_path)
            # Try to extract text from supported file types
            try:
                if file_name.lower().endswith('.pdf'):
                    import fitz
                    with fitz.open(file_path) as pdf:
                        file_text = '\n'.join([page.get_text() for page in pdf])
                elif file_name.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(file_path)
                    file_text = '\n'.join([p.text for p in doc.paragraphs])
                elif file_name.lower().endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        file_text = f.read()
                elif file_name.lower().endswith('.csv'):
                    import pandas as pd
                    df = pd.read_csv(file_path)
                    file_text = df.to_string()
                elif file_name.lower().endswith('.xlsx'):
                    import pandas as pd
                    df = pd.read_excel(file_path)
                    file_text = df.to_string()
                elif file_name.lower().endswith(('.png', '.jpg', '.jpeg')):
                    # OCR for images containing text/tables
                    try:
                        import cv2
                        import pytesseract
                        from PIL import Image
                        import numpy as np
                        
                        # Read image and convert to OpenCV format
                        pil_image = Image.open(file_path)
                        cv_image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
                        
                        # Preprocess image for better OCR
                        gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
                        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                        
                        # Extract text using Tesseract OCR
                        file_text = pytesseract.image_to_string(thresh, config='--psm 6')
                        
                        if not file_text.strip():
                            # Try with different PSM mode for tables
                            file_text = pytesseract.image_to_string(thresh, config='--psm 3')
                        
                        if not file_text.strip():
                            file_text = "No text could be extracted from this image."
                    except ImportError:
                        file_text = "OCR libraries not installed. Please install opencv-python, pytesseract, and pillow."
                    except Exception as e:
                        file_text = f"OCR error: {str(e)}"
                # Add more file types as needed
            except Exception as e:
                file_text = f"[Could not extract text from file: {e}]"
        if user_text or file_name:
            chat_history.append({'role': 'user', 'text': user_text, 'file_name': file_name})
            # Gemini AI response (chat key)
            if GEMINI_API_KEY_CHAT:
                prompt = user_text or ''
                if file_text:
                    prompt = f"File content:\n{file_text}\n\nUser question: {user_text}"
                try:
                    headers = {'Content-Type': 'application/json'}
                    params = {'key': GEMINI_API_KEY_CHAT}
                    data = {
                        "contents": [{"parts": [{"text": prompt}]}]
                    }
                    response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data)
                    if response.status_code == 200:
                        gemini_data = response.json()
                        bot_response = gemini_data['candidates'][0]['content']['parts'][0]['text']
                    else:
                        bot_response = f"[Gemini API error: {response.text}]"
                except Exception as e:
                    bot_response = f"[Gemini error: {e}]"
            else:
                bot_response = f"You said: {user_text} (Gemini AI not configured)"
                if file_name:
                    bot_response += f" (File received: {file_name})"
            chat_history.append({'role': 'bot', 'text': bot_response, 'file_name': None})
            session['chat_history'] = chat_history
    return render_template('internal_chatbot.html', chat_history=chat_history, error=error)

@prompt_optimizer_bp.route('/', methods=['GET', 'POST'])
def prompt_optimizer_home():
    optimized = None
    explanation = None
    error = None
    original_prompt = ""
    selected_tone = "neutral"
    selected_role = "general"

    if request.method == 'POST':
        original_prompt = request.form.get('prompt')
        selected_tone = request.form.get('tone') or 'neutral'
        selected_role = request.form.get('role') or 'general'

        if not original_prompt.strip():
            error = 'Prompt is required.'
        elif GEMINI_API_KEY_CHAT:
            try:
                headers = {'Content-Type': 'application/json'}
                params = {'key': GEMINI_API_KEY_CHAT}

                gemini_prompt = (
                    f"You are a helpful AI assistant optimizing prompts for LLMs like ChatGPT and Gemini.\n\n"
                    f"User Role: {selected_role.capitalize()}\n"
                    f"Tone: {selected_tone.capitalize()}\n\n"
                    f"Task:\nImprove the following prompt to be clearer, more detailed, and contextually appropriate.\n"
                    f"Then explain briefly what was improved.\n\n"
                    f"Original Prompt: {original_prompt}"
                )

                data = {
                    "contents": [
                        {
                            "parts": [
                                {"text": gemini_prompt}
                            ]
                        }
                    ]
                }

                response = requests.post(GEMINI_API_URL, headers=headers, params=params, json=data)

                if response.status_code == 200:
                    gemini_data = response.json()
                    raw_text = gemini_data['candidates'][0]['content']['parts'][0]['text']
                    
                    # Optional: Split explanation if Gemini separates it
                    if "Explanation:" in raw_text:
                        optimized, explanation = raw_text.split("Explanation:", 1)
                    else:
                        optimized = raw_text
                else:
                    error = f"[Gemini API error: {response.text}]"
            except Exception as e:
                error = f"[Gemini error: {e}]"
        else:
            error = 'Gemini API key not configured.'

    return render_template(
        'prompt_optimizer.html',
        optimized=optimized,
        explanation=explanation,
        error=error,
        original_prompt=original_prompt,
        selected_tone=selected_tone,
        selected_role=selected_role
    )

@chat_history_bp.route('/', methods=['GET'])
def chat_history_home():
    history = session.get('history', [])
    return render_template('chat_history.html', history=history)

# Register blueprints AFTER all routes are defined (text_to_image bp REMOVED)
app.register_blueprint(document_assistant_bp)
app.register_blueprint(developer_copilot_bp)
app.register_blueprint(internal_chatbot_bp)
app.register_blueprint(chat_history_bp)
app.register_blueprint(prompt_optimizer_bp)

# Helper function for basic document processing
def process_document_basic(file_text, query_type, query):
    if query_type == 'summary':
        sentences = file_text.split('.')
        summary = '. '.join(sentences[:3]) + '.'
        if len(summary) > 300:
            summary = summary[:300] + '...'
        return f"Document Summary:\n{summary}\n\nDocument Length: {len(file_text)} characters"
    elif query_type == 'extract':
        lines = file_text.split('\n')
        key_lines = [line.strip() for line in lines if line.strip() and len(line.strip()) > 20]
        return f"Key Content (first 10 lines):\n" + '\n'.join(key_lines[:10])
    elif query_type == 'context':
        if query:
            query_lower = query.lower()
            file_lower = file_text.lower()
            if query_lower in file_lower:
                start = file_lower.find(query_lower)
                context = file_text[max(0, start-100):start+len(query)+100]
                return f"Context for '{query}':\n...{context}..."
            else:
                return f"Query '{query}' not found in document."
        else:
            return "Please enter a search term to find context."
    return "Unknown query type."

def add_to_history(module, user_input, response):
    if 'history' not in session:
        session['history'] = []
    session['history'].append({
        'module': module,
        'user_input': user_input,
        'response': response
    })
    session.modified = True

@app.route('/')
def home():
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=False, host='0.0.0.0')

